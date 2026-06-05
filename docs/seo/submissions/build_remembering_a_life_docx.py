from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "docs/seo/submissions/remembering-a-life-how-to-turn-an-obituary-into-a-life-story.md"
OUTPUT = ROOT / "docs/seo/submissions/remembering-a-life-how-to-turn-an-obituary-into-a-life-story.docx"


def set_font(run, *, size=11, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def paragraph_format(paragraph, *, before=0, after=8, line_spacing=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text):
    url = "https://www.finalnotes.page/how-to-write-an-obituary-story/"
    if url not in text:
        run = paragraph.add_run(text)
        set_font(run)
        return
    before, after = text.split(url, 1)
    if before:
        run = paragraph.add_run(before)
        set_font(run)
    add_hyperlink(paragraph, url, url)
    if after:
        run = paragraph.add_run(after)
        set_font(run)


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph_format(paragraph, before=18 if level == 1 else 14, after=6)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 14, bold=False)
    return paragraph


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph_format(paragraph)
    add_text_with_links(paragraph, text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph_format(paragraph, after=4)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def build_docx():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)
    for style_name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 14)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False

    title = lines[0].removeprefix("# ").strip()
    title_p = doc.add_paragraph()
    paragraph_format(title_p, after=3)
    title_run = title_p.add_run(title)
    set_font(title_run, size=26, bold=False)

    byline = doc.add_paragraph()
    paragraph_format(byline, after=12)
    byline_run = byline.add_run(lines[2].strip())
    set_font(byline_run, size=11, italic=True)

    in_bullets = False
    for raw in lines[3:]:
        line = raw.strip()
        if not line:
            in_bullets = False
            continue
        if line.startswith("## "):
            add_heading(doc, line.removeprefix("## "), 2)
            continue
        if line.startswith("- "):
            in_bullets = True
            add_bullet(doc, line.removeprefix("- "))
            continue
        if line.startswith("# "):
            continue
        if line == "## Author bio":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_body_paragraph(doc, line)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
